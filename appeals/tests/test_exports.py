from datetime import datetime
from io import BytesIO

import pytest
from django.utils import timezone
from docx import Document
from openpyxl import load_workbook

from appeals.exports import (
    REPORT_AUTHOR,
    report_to_docx,
    report_to_xlsx,
)
from appeals.reports import AppealReport, Breakdown


def _sample_report():
    return AppealReport(
        total=5,
        overdue=2,
        by_status=[
            Breakdown(label="Новые", count=1),
            Breakdown(label="В работе", count=1),
            Breakdown(label="Закрытые", count=3),
        ],
        by_category=[
            Breakdown(label="Справки", count=3),
            Breakdown(label="Переводы", count=2),
        ],
        by_department=[Breakdown(label="Учебный отдел", count=5)],
    )


def _fixed_moment():
    # Фиксированный момент времени, чтобы проверять форматирование без now().
    return timezone.make_aware(datetime(2026, 6, 2, 14, 30))


@pytest.mark.unit
def test_xlsx_export_has_all_sheets_and_summary():
    content = report_to_xlsx(_sample_report(), generated_at=_fixed_moment())

    workbook = load_workbook(BytesIO(content))
    assert workbook.sheetnames == [
        "Сводка",
        "По статусам",
        "По категориям",
        "По отделам",
    ]

    summary = workbook["Сводка"]
    text = "\n".join(
        str(cell.value) for row in summary.iter_rows() for cell in row if cell.value is not None
    )
    assert "Всего обращений" in text
    assert "02.06.2026 14:30" in text
    assert REPORT_AUTHOR in text


@pytest.mark.unit
def test_xlsx_export_writes_summary_values():
    content = report_to_xlsx(_sample_report(), generated_at=_fixed_moment())

    summary = load_workbook(BytesIO(content))["Сводка"]
    values = {
        summary.cell(row=row, column=1).value: summary.cell(row=row, column=2).value
        for row in range(1, summary.max_row + 1)
    }
    assert values["Всего обращений"] == 5
    assert values["Просрочено"] == 2


@pytest.mark.unit
def test_xlsx_breakdown_sheet_lists_rows():
    content = report_to_xlsx(_sample_report(), generated_at=_fixed_moment())

    sheet = load_workbook(BytesIO(content))["По категориям"]
    rows = [
        (sheet.cell(row=row, column=1).value, sheet.cell(row=row, column=2).value)
        for row in range(2, sheet.max_row + 1)
    ]
    assert rows == [("Справки", 3), ("Переводы", 2)]


@pytest.mark.unit
def test_xlsx_export_returns_nonempty_bytes():
    content = report_to_xlsx(_sample_report(), generated_at=_fixed_moment())

    assert isinstance(content, bytes)
    assert content


@pytest.mark.unit
def test_docx_export_contains_summary_text():
    content = report_to_docx(_sample_report(), generated_at=_fixed_moment())

    document = Document(BytesIO(content))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Сводный отчёт по обращениям" in text
    assert "Всего обращений: 5" in text
    assert "Просрочено: 2" in text
    assert "02.06.2026 14:30" in text
    assert REPORT_AUTHOR in text


@pytest.mark.unit
def test_docx_export_renders_breakdown_tables():
    content = report_to_docx(_sample_report(), generated_at=_fixed_moment())

    document = Document(BytesIO(content))
    # Три разбивки → три таблицы.
    assert len(document.tables) == 3

    category_table = document.tables[1]
    body_rows = [(row.cells[0].text, row.cells[1].text) for row in category_table.rows[1:]]
    assert body_rows == [("Справки", "3"), ("Переводы", "2")]


@pytest.mark.unit
def test_docx_export_handles_empty_breakdown():
    empty = AppealReport(total=0, overdue=0)
    content = report_to_docx(empty, generated_at=_fixed_moment())

    document = Document(BytesIO(content))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    # Пустые разбивки заменяются заметкой, таблиц нет.
    assert "Нет данных." in text
    assert document.tables == []
