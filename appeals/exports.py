from datetime import datetime
from io import BytesIO

from django.utils import timezone
from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Font

from appeals.reports import AppealReport, Breakdown

# Общие подписи документа — те же название и автор, что и в подвале сайта,
# чтобы выгрузка выглядела частью того же сервиса.
REPORT_TITLE = "Журнал учета телефонных обращений обучающихся"
REPORT_SUBTITLE = "Сводный отчёт по обращениям"
REPORT_AUTHOR = "Автор: Лавров Дмитрий Андреевич"


def report_to_xlsx(report: AppealReport, *, generated_at: datetime) -> bytes:
    """Формирует .xlsx-файл сводки и возвращает его содержимое байтами.

    Каждая разбивка выносится на отдельный лист, плюс лист со сводными
    показателями. Время формирования передаётся снаружи, чтобы функция
    оставалась чистой и предсказуемой в тестах.
    """
    workbook = Workbook()

    summary = workbook.active
    summary.title = "Сводка"
    _write_xlsx_summary(summary, report, generated_at=generated_at)

    _write_xlsx_breakdown(workbook.create_sheet("По статусам"), report.by_status)
    _write_xlsx_breakdown(workbook.create_sheet("По категориям"), report.by_category)
    _write_xlsx_breakdown(workbook.create_sheet("По отделам"), report.by_department)

    buffer = BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def report_to_docx(report: AppealReport, *, generated_at: datetime) -> bytes:
    """Формирует .docx-файл сводки и возвращает его содержимое байтами.

    Сводные показатели идут списком, а каждая разбивка — отдельной таблицей.
    Время формирования передаётся снаружи, как и для xlsx-выгрузки.
    """
    document = Document()
    document.add_heading(REPORT_TITLE, level=0)
    document.add_heading(REPORT_SUBTITLE, level=1)

    meta = document.add_paragraph()
    meta.add_run(f"Сформирован: {_format_moment(generated_at)}\n")
    meta.add_run(REPORT_AUTHOR)

    document.add_heading("Показатели", level=2)
    document.add_paragraph(f"Всего обращений: {report.total}")
    document.add_paragraph(f"Просрочено: {report.overdue}")

    _write_docx_table(document, "По статусам", report.by_status)
    _write_docx_table(document, "По категориям", report.by_category)
    _write_docx_table(document, "По отделам", report.by_department)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _write_xlsx_summary(
    sheet,
    report: AppealReport,
    *,
    generated_at: datetime,
) -> None:
    bold = Font(bold=True)

    title_cell = sheet.cell(row=1, column=1, value=REPORT_SUBTITLE)
    title_cell.font = bold
    sheet.cell(row=2, column=1, value=f"Сформирован: {_format_moment(generated_at)}")
    sheet.cell(row=3, column=1, value=REPORT_AUTHOR)

    sheet.cell(row=5, column=1, value="Показатель").font = bold
    sheet.cell(row=5, column=2, value="Значение").font = bold
    sheet.cell(row=6, column=1, value="Всего обращений")
    sheet.cell(row=6, column=2, value=report.total)
    sheet.cell(row=7, column=1, value="Просрочено")
    sheet.cell(row=7, column=2, value=report.overdue)

    sheet.column_dimensions["A"].width = 28
    sheet.column_dimensions["B"].width = 16


def _write_xlsx_breakdown(sheet, rows: list[Breakdown]) -> None:
    bold = Font(bold=True)
    sheet.cell(row=1, column=1, value="Группа").font = bold
    sheet.cell(row=1, column=2, value="Обращений").font = bold
    for index, row in enumerate(rows, start=2):
        sheet.cell(row=index, column=1, value=row.label)
        sheet.cell(row=index, column=2, value=row.count)

    sheet.column_dimensions["A"].width = 36
    sheet.column_dimensions["B"].width = 14


def _write_docx_table(document: Document, heading: str, rows: list[Breakdown]) -> None:
    document.add_heading(heading, level=2)
    if not rows:
        document.add_paragraph("Нет данных.")
        return

    table = document.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    header = table.rows[0].cells
    header[0].text = "Группа"
    header[1].text = "Обращений"
    for cell in header:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)

    for row in rows:
        cells = table.add_row().cells
        cells[0].text = row.label
        cells[1].text = str(row.count)


def _format_moment(moment: datetime) -> str:
    # Показываем время в локальной зоне проекта в привычном формате ДД.ММ.ГГГГ.
    return timezone.localtime(moment).strftime("%d.%m.%Y %H:%M")
